#!/usr/bin/env python3
"""Build VioNER_Article.docx — journal-manuscript version of the thesis.

Produces a single-column, Times New Roman, A4 manuscript with numbered
sections, captioned tables, and numbered equations, suitable for
submission to an open-access journal (IEEE Access / PLOS ONE style).

Run with an interpreter that has python-docx installed.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "VioNER_Article.docx")

doc = Document()

# ---------------- page + base styles ----------------
for section in doc.sections:
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)  # A4
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(2.54)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, italic in (("Heading 1", 14, False), ("Heading 2", 12, False), ("Heading 3", 12, True)):
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.italic = italic
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def para(text, style=None, align=None, bold=False, italic=False, size=None, space_after=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(text, level):
    doc.add_heading(text, level=level)


def equation(text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True
    p.add_run("        (%d)" % number)


TABLE_N = [0]


def add_table(caption, headers, rows, col_widths=None):
    TABLE_N[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    rb = cap.add_run("Table %d. " % TABLE_N[0])
    rb.bold = True
    rb.font.size = Pt(11)
    rc = cap.add_run(caption)
    rc.font.size = Pt(11)

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E7E6E6")
        cell._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    para("", space_after=4)
    return t


# ================= TITLE PAGE BLOCK =================
para("Knowledge-Grounded Named Entity Recognition for Extracting Structured "
     "Violent-Event Information from African Conflict Reporting",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=14)

para("Binalfew Kassa¹*  [add advisor/co-author here]",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
para("¹ Department of Computer Science, Addis Ababa University, Addis Ababa, Ethiopia",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=2)
para("* Corresponding author: binalfew@yahoo.com",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=18)

# ================= ABSTRACT =================
heading("Abstract", 1)
para(
    "Quantitative conflict research and operational early warning both run on structured "
    "event records, yet nearly all such records for Africa are still produced by analysts "
    "reading news reports one at a time. This paper presents an approach to named entity "
    "recognition (NER) for extracting who-did-what-to-whom information from English-language "
    "reporting of violent events in Africa. Three design decisions distinguish the approach. "
    "First, the entity schema is restricted to eight types — actor, victim, action, date, "
    "region, city, district, and casualties — each verified in a pilot annotation study to be "
    "reliably grounded in source text; categories that annotators could only infer, such as event "
    "type and country, are recovered by deterministic post-processing instead of being learned. "
    "Second, severe class imbalance — 78% of tokens carry no entity label, and the "
    "operationally critical victim and casualty classes account for under 3% of entity tokens "
    "each — is addressed by fine-tuning BERT under focal loss combined with inverse-frequency "
    "class weighting, on a 50,000-example corpus derived from ACLED event descriptions through "
    "stratified diversity sampling and template augmentation. Third, a curated knowledge base of "
    "African armed groups and conflict-affected locations validates and enriches the model "
    "output, and a four-level taxonomy of African violent events with roughly ninety-five "
    "terminal categories classifies each extracted record. On a held-out validation set the "
    "model reaches macro F1 of 0.887 and micro F1 of 0.909. A controlled ablation shows the two "
    "imbalance treatments are complementary: together they lift the rarest entity (victim) by "
    "eleven F1 points over plain cross-entropy without degrading any frequent entity. The "
    "knowledge base canonicalises 64.3% of high-confidence actor mentions and flags "
    "geographically implausible extractions for analyst review. The taxonomy, annotation "
    "guidelines, and knowledge base are released publicly.")
para("Keywords: named entity recognition; event extraction; conflict monitoring; class "
     "imbalance; focal loss; BERT; knowledge bases; African NLP", italic=True, space_after=12)

# ================= 1 INTRODUCTION =================
heading("1. Introduction", 1)
para(
    "The quantitative study of armed conflict, and the early-warning practice built on top of "
    "it, depends on structured event records: who acted, who was harmed, where and when it "
    "happened, and at what human cost. The databases that supply those records have proved "
    "their value many times over — the Armed Conflict Location and Event Data Project "
    "(ACLED) [1] and the Uppsala Conflict Data Program (UCDP) [2] underpin a large share of "
    "contemporary conflict scholarship — but they are produced substantially by human "
    "coders reading news reports one article at a time. Fully automated alternatives such as "
    "GDELT [3] trade precision for scale, and their output has not displaced hand-coded data "
    "in analytical practice. Meanwhile the volume of reporting keeps growing. Monitoring "
    "organisations such as the African Union's Continental Early Warning System receive a "
    "daily inflow of news that exceeds what their analyst rosters can read, let alone convert "
    "into structured form.")
para(
    "Manual coding suffers from two distinct problems that compound each other. The first is "
    "throughput: the share of relevant reporting that ever becomes a structured record is "
    "capped by the number of hours analysts can spend reading. The second is consistency: two "
    "coders given the same article will sometimes diverge on basic judgements — whether an "
    "incident is a battle or violence against civilians, whether a place name denotes a city "
    "or a district, to whom a fatality count should be attributed. Hiring more coders eases "
    "the first problem while aggravating the second; tightening the coding manual does the "
    "reverse. A learned extractor is attractive precisely because it presses on both at once: "
    "it applies one consistent judgement to every article, and it can read the entire inflow.")
para(
    "What makes the problem tractable now is transfer learning. Pre-trained transformer "
    "encoders [4], and BERT in particular [5], allow a token-classification model to be "
    "fine-tuned to competitive accuracy from a corpus of tens of thousands of examples rather "
    "than millions. Yet African conflict text is exactly the setting where off-the-shelf "
    "models stumble. The actors that matter — Boko Haram, JNIM, the M23, the Rapid Support "
    "Forces, Fano, Anti-Balaka — are poorly represented in generic pre-training corpora, "
    "and the MasakhaNER benchmark has documented that widely used multilingual encoders "
    "underperform on African text relative to encoders exposed to African corpora during "
    "pre-training [6]. Domain fine-tuning is therefore not optional in this setting; it is the "
    "core of the method.")
para(
    "Three obstacles are specific to this domain and shape the contributions of this paper. "
    "The first concerns what can legitimately be supervised. Several semantic categories an "
    "analyst cares about — the type of an event, the country it occurred in — are "
    "frequently never stated in the text at all but inferred by the reader, and a sequence "
    "labeller trained on labels that cannot be located in its input is being trained on noise. "
    "The second is class imbalance of a severity that ordinary cross-entropy training handles "
    "badly: in the corpus assembled here, roughly 78% of tokens fall outside any entity, and "
    "the rarest entity classes — the victims and casualty counts an analyst most needs "
    "— each account for under 3% of the entity tokens that remain. The third is trust: an "
    "extraction system whose output cannot be audited against domain knowledge adds work for "
    "an analyst instead of removing it.")
para(
    "This paper describes VioNER, an extraction approach built around those three obstacles, "
    "and makes four contributions. (1) A grounded eight-entity schema for violent-event NER, "
    "selected by a pilot annotation study that measured, per candidate entity type, how often "
    "its value can be located verbatim in source text; the schema's inclusion and exclusion "
    "rules are documented and released. (2) An imbalance-aware fine-tuning recipe that "
    "combines focal loss with inverse-frequency class weighting, supported by a controlled "
    "ablation showing the two ingredients are complementary rather than redundant. (3) A "
    "four-level taxonomy of African violent events with roughly ninety-five terminal "
    "categories, synthesised from ACLED, UCDP, and the PMVE ontology with African-specific "
    "extensions, together with a curated knowledge base of about 150 armed groups and 200 "
    "conflict-affected cities; both artefacts are released publicly. (4) An empirical "
    "evaluation comprising per-entity analysis, measurement of the knowledge-base layer's "
    "operational impact, and a manual error analysis of 300 misclassified events.")
para(
    "The remainder of the paper is organised as follows. Section 2 reviews related work. "
    "Section 3 presents the methodology: the grounded schema, the taxonomy, corpus "
    "construction, the training objective, and knowledge-base validation. Section 4 describes "
    "the experimental setup, Section 5 reports results, Section 6 discusses findings and "
    "limitations, and Section 7 concludes.")

# ================= 2 RELATED WORK =================
heading("2. Related Work", 1)
heading("2.1 Event extraction from news", 2)
para(
    "Supervised event extraction descends largely from the Automatic Content Extraction "
    "programme [7], which formalised events as typed predicates with arguments, and from "
    "Ahn's decomposition of the task into trigger and argument identification [8]. An "
    "alternative to any fixed event typology is the journalistic 5W1H frame — who, what, "
    "whom, where, when, how — which has also been operationalised as a machine reading "
    "comprehension task [9]. The frame suits news text unusually well, since reporters are "
    "trained to answer precisely those questions early in an article, and it has the useful "
    "property of decoupling what is extracted from which event types are assumed to exist. "
    "That is the frame adopted here: event-type assignment happens after extraction, against "
    "an explicit taxonomy, and never enters the supervised learning problem.")
para(
    "Surveying the field, Hogenboom et al. distinguish data-driven, knowledge-driven, and "
    "hybrid extraction methods, and observe that hybrid pipelines dominate operational "
    "deployments because domain knowledge is encoded where rules express it naturally while "
    "statistical learning covers what rules cannot [10]. The system described here is hybrid "
    "in exactly that sense. Earlier operational systems illustrate the trade-off. Tanev et "
    "al. built a real-time multilingual pipeline for global crisis monitoring on hand-written "
    "patterns [11], and the NEXUS system extracted violent events against the Politically "
    "Motivated Violent Events (PMVE) ontology, likewise with linguistic patterns [12]. "
    "Pattern-based recognition, however, scales poorly across vocabulary shifts: porting a "
    "pattern base tuned to European security incidents onto a continent with several hundred "
    "active armed groups and tens of thousands of locality names would mean rebuilding it. A "
    "fine-tuned encoder amortises that vocabulary work into pre-training.")
heading("2.2 Named entity recognition", 2)
para(
    "NER methodology has moved through three eras: feature-engineered sequence models "
    "culminating in conditional random fields [13], neural architectures built on "
    "bidirectional LSTMs with CRF decoding [14], [15], and fine-tuned transformer encoders, "
    "for which the recipe — contextual encoder, linear classification head, cross-entropy "
    "or focal loss — has been stable since BERT [5]. The property that matters for a "
    "low-resource domain is transfer: a 50,000-example specialised corpus is far too small to "
    "train a sequence model from scratch but is sufficient to fine-tune a pre-trained "
    "encoder. For African settings specifically, MasakhaNER established NER benchmarks for "
    "ten African languages over a four-type schema (person, organisation, location, date) "
    "[6], and subsequent encoders such as AfroLM [16] and XLM-RoBERTa [17] target African or "
    "multilingual text directly. The present work differs in aim: its input is "
    "English-language reporting of African events rather than African-language text, and its "
    "schema is task-specific and richer, distinguishing the semantic roles — perpetrator "
    "versus victim, region versus district — that a conflict analyst actually uses.")
heading("2.3 Class imbalance in token classification", 2)
para(
    "BIO-encoded NER produces imbalance by construction, since most tokens in ordinary prose "
    "belong to no entity. Three families of remedy exist. Resampling [18] is complicated in "
    "token classification by a unit mismatch: oversampling is done per sentence while the "
    "imbalance lives at the token level, so each sentence recruited for a rare entity drags "
    "in dozens of majority-class tokens. Cost-sensitive losses reweight classes, classically "
    "by inverse frequency [19] or by effective sample number [20]. Focal loss [21] instead "
    "down-weights examples the model already classifies confidently, concentrating gradient "
    "on hard cases; it was proposed for dense object detection, where the "
    "foreground–background ratio is even more extreme than in NER. The two mechanisms act "
    "on different axes — one on class frequency, the other on example difficulty — "
    "and the ablation in Section 5.3 confirms empirically that they compose.")
heading("2.4 Conflict event data, taxonomies, and prior African-context work", 2)
para(
    "ACLED's own coding scheme is a two-level taxonomy of six event types and roughly "
    "twenty-five sub-event types [1]; UCDP distinguishes state-based conflict, non-state "
    "conflict, and one-sided violence [2]; GDELT codes events against the broad CAMEO "
    "hierarchy [3], [22]. None of these reaches the granularity an African-focused analyst "
    "needs for categories such as pastoralist–farmer clashes or communal cattle raiding, "
    "which the frameworks fold into broader buckets. On the extraction side, the closest "
    "predecessor to this work extracted 5W attributes of African violent events using "
    "Stanford CoreNLP and Weka classifiers over a modest annotated corpus [23], establishing "
    "feasibility but predating transformer fine-tuning; related 5W1H extraction with "
    "ontology population has been demonstrated for Chinese news [24]. Hierarchical "
    "classification methodology is surveyed in [25]. Against this background, the present "
    "work contributes the first transformer-based NER model targeted specifically at African "
    "violent-event extraction at the 50,000-example scale, a public four-level taxonomy of "
    "African violent events, and a knowledge-base validation layer coupling a learned "
    "extractor with curated domain reference data.")

# ================= 3 METHODOLOGY =================
heading("3. Methodology", 1)
para(
    "The extraction pipeline is deliberately modular and hybrid. A fine-tuned BERT model "
    "performs token-level entity recognition; deterministic post-processing assembles token "
    "predictions into spans with confidence scores, filters them against per-category "
    "thresholds, validates and enriches them against a curated knowledge base, groups them "
    "into 5W1H slots, and assigns each event a position in a hierarchical taxonomy. "
    "Statistical learning is confined to the one stage that needs generalisation over "
    "surface forms; every stage whose job is look-up or rule application is implemented "
    "deterministically, which keeps each step inspectable and separately testable. The "
    "subsections below describe the schema, the taxonomy, the corpus, the training "
    "objective, and the knowledge-base layer in turn.")
heading("3.1 A grounded entity schema", 2)
para(
    "The schema was fixed by a grounding rule: an entity type enters the supervised schema "
    "only if a human annotator can locate its value verbatim in the source text on a "
    "reliable majority of occurrences. The rule was enforced through a pilot annotation "
    "study on a sample of conflict articles, and it eliminated two categories that an "
    "earlier candidate schema had included. Event type failed because reports rarely name "
    "the category of an incident — whether a given attack is an ambush or a raid is "
    "usually the reader's inference, not the writer's words — and its grounding rate "
    "came in below 60%. Country failed for the opposite reason: writers omit the country "
    "when a city or region name already implies it. Both categories are recovered "
    "downstream — event type by the taxonomy classifier from the action verb and "
    "context, country by knowledge-base look-up from the most specific extracted location "
    "— so nothing is lost from the final record. What remains is the eight-entity "
    "schema of Table 1, encoded in BIO format, which with the outside label yields "
    "seventeen output classes.")
add_table(
    "The eight-entity grounded schema, organised by 5W1H role.",
    ["Role", "Entity", "Description and examples"],
    [
        ["WHO", "ACTOR", "Armed groups, organisations, state forces (“Boko Haram”, “M23 rebels”, “gunmen”)"],
        ["WHOM", "VICTIM", "Persons or groups affected (“civilians”, “villagers”, “protesters”)"],
        ["WHAT", "ACTION", "Verbs describing the event (“attacked”, “killed”, “ambushed”)"],
        ["WHEN", "DATE", "Temporal expressions (“on Monday”, “January 15, 2024”)"],
        ["WHERE", "REGION", "States, provinces, regions (“Borno State”, “North Kivu”)"],
        ["WHERE", "CITY", "Cities, towns, villages (“Maiduguri”, “Goma”, “Mogadishu”)"],
        ["WHERE", "DISTRICT", "Districts, counties, localities (“Bama”, “Masisi”)"],
        ["HOW", "CASUALTIES", "Death and injury counts (“killed 12”, “3 dead”, “5 injured”)"],
    ],
    col_widths=[2.2, 3.0, 10.7],
)
heading("3.2 A hierarchical taxonomy of African violent events", 2)
para(
    "Extracted events are classified against a four-level taxonomy developed for this work. "
    "Level 1 contains four broad categories; Level 2 introduces eighteen intermediate types; "
    "Level 3 refines these into roughly fifty specific event types; and Level 4 adds some "
    "twenty detailed subtypes where operational distinctions warrant them, for a total of "
    "approximately ninety-five terminal categories. The synthesis draws on ACLED's event "
    "types [1], UCDP's distinction between state-based, non-state, and one-sided violence "
    "[2], and the violence classes of the PMVE ontology [12], extended with categories that "
    "reflect African conflict dynamics — pastoralist–farmer clashes, communal cattle "
    "raiding, election violence — which none of those frameworks covers at comparable "
    "depth. Every terminal category carries a definition, classification criteria, typical "
    "lexical cues, and worked examples, and every pair of overlapping categories carries an "
    "explicit decision rule, so classification never falls back on annotator intuition. "
    "Table 2 summarises the upper two levels.")
add_table(
    "Upper two levels of the violent-event taxonomy.",
    ["Level 1 category", "Level 2 types"],
    [
        ["Political Violence", "Rebellion / Armed Insurgency; Terrorism; Coup and Regime-Change Violence; Election Violence; Political Repression"],
        ["Criminal Violence", "Organised-Crime Violence; Armed Robbery / Banditry; Kidnapping for Ransom; Criminal Gang Violence"],
        ["Communal Violence", "Ethnic / Tribal Conflict; Religious Violence; Resource-Based Conflict; Pastoralist–Farmer Clashes"],
        ["State Violence Against Civilians", "Extrajudicial Killings; Repression of Protests; Mass Atrocities by State Forces; Forced Displacement; Arbitrary Detention with Violence"],
    ],
    col_widths=[4.6, 11.3],
)
para(
    "Within the pipeline, taxonomy assignment is a deterministic post-extraction step: a "
    "rule-based classifier inspects the actor type, the action verb, and contextual cues "
    "such as weapons and casualty patterns, consults the knowledge base for the actor's "
    "known classification, and descends the tree as far as the evidence supports. An "
    "analyst can stop at whatever level of granularity they trust, which is the practical "
    "argument for a deep taxonomy despite the difficulty of its finest distinctions.")
heading("3.3 Corpus construction", 2)
para(
    "The training corpus derives from the ACLED open-data export of 212,590 African event "
    "records. Each record pairs structured columns (actors, location, administrative "
    "region, date, fatalities) with a free-text note describing the event; projecting the "
    "column values onto their occurrences in the tokenised note yields BIO labels without "
    "manual annotation. Two corpus-level problems then had to be addressed. The first is "
    "redundancy: ACLED notes follow house conventions, and a large share of records are "
    "near-paraphrases of the same sentence pattern, so a model trained on the full extract "
    "plateaus early — an effect confirmed empirically when an initial model trained on "
    "all 212,000 records scored lower on rare-entity F1 than later models trained on "
    "smaller, more diverse subsets. A stratified diversity sampler therefore selects a "
    "35,000-example subset that over-represents sentences containing rare entity types "
    "while maximising lexical and structural diversity within each stratum. The sampling "
    "unit is the sentence but the target of the stratification is the entity type, which "
    "limits the extent to which recruiting rare entities also recruits majority-class "
    "tokens.")
para(
    "The second problem is vocabulary coverage. ACLED's editorial style under-represents "
    "constructions common in raw news prose, notably active-voice action verbs and "
    "descriptive victim phrases. Template-based augmentation generates a further 15,000 "
    "synthetic examples from verb lexicons and slot-filling patterns aimed at exactly those "
    "gaps. The combined 50,000-example corpus — roughly 30% synthetic — is split "
    "80/20 into training and validation partitions, stratified by entity-type presence. "
    "The implications of evaluating on a partition that includes synthetic examples are "
    "taken up in Section 6.")
heading("3.4 Model and training objective", 2)
para(
    "The extractor is bert-base-cased with a linear token-classification head over the "
    "seventeen BIO labels. The cased variant is a deliberate choice: English conflict "
    "reporting marks actor and place names with capitalisation, and discarding case "
    "discards signal. WordPiece sub-word labels follow the standard convention — the "
    "first sub-word of a word inherits the word's label, subsequent sub-words inherit it "
    "with any leading B- prefix rewritten to I-, and special tokens receive the ignore "
    "index so the loss skips them.")
para(
    "The training objective combines focal loss with inverse-frequency class weighting. "
    "For a token with true class y and predicted probability pᵧ, the loss is")
equation("FL(p, y) = −αᵧ · (1 − pᵧ)^γ · log pᵧ", 1)
para(
    "where γ ≥ 0 controls how strongly already-confident predictions are discounted "
    "(γ = 0 recovers cross-entropy) and αᵧ is a per-class weight. Weights are "
    "computed once from the training distribution as")
equation("αᴄ = T / (C · max(fᴄ, 1))", 2)
para(
    "with T the total token count, C the number of classes, and fᴄ the frequency of "
    "class c, clipped at a maximum of 10 to prevent gradient instability. Under this "
    "scheme the outside label receives a weight of 0.07 while the three rarest entity "
    "classes (victim, action, casualties) saturate at the cap. Label smoothing of 0.1 is "
    "applied as mild regularisation. Training follows the standard BERT fine-tuning "
    "recipe (Table 3); the focusing parameter γ = 2 follows the original focal-loss "
    "recommendation [21], and no hyperparameter search was performed beyond it. Early "
    "stopping monitors validation loss.")
add_table(
    "Training configuration.",
    ["Setting", "Value"],
    [
        ["Pre-trained encoder", "bert-base-cased (12 layers, 110M parameters)"],
        ["Output classes", "17 BIO labels"],
        ["Maximum sequence length", "512"],
        ["Effective batch size", "32 (16 × 2 gradient accumulation)"],
        ["Optimiser / learning rate", "AdamW, 2 × 10⁻⁵, weight decay 0.01"],
        ["Schedule", "500 warm-up steps; ReduceLROnPlateau (factor 0.5, patience 2)"],
        ["Loss", "Focal (γ = 2) + inverse-frequency weights (cap 10) + smoothing 0.1"],
        ["Early stopping", "patience 5, threshold 0.001 on validation loss"],
        ["Maximum epochs", "10"],
    ],
    col_widths=[6.0, 9.9],
)
heading("3.5 Knowledge-base validation and post-processing", 2)
para(
    "At inference time, contiguous B-/I- runs are assembled into spans, each carrying the "
    "mean of its sub-token softmax probabilities as a confidence score. Spans below a "
    "per-category threshold are discarded; thresholds (0.60 for actions, 0.70 for actors, "
    "victims, and locations, 0.75 for casualties, 0.80 for dates) were set by inspecting "
    "validation-set errors rather than by formula, on the reasoning that the cost of a "
    "false positive differs by category — a wrong date is glaring and cheap to filter "
    "aggressively, while a low-confidence action verb is usually propped up by context. "
    "Long documents that describe several incidents are first segmented into event-bearing "
    "sentence groups using discourse cues (“in a separate incident”, "
    "“meanwhile”), so that one record is produced per incident rather than one "
    "blended record per article.")
para(
    "Surviving spans are then grounded against a curated knowledge base (Table 4) holding "
    "roughly 150 armed groups with aliases, country of operation, region, and type; about "
    "200 conflict-affected cities mapped to country and parent administrative unit; all 54 "
    "African countries; and a categorised weapons lexicon. The knowledge base serves two "
    "distinct functions. Enrichment attaches canonical identifiers and metadata: an actor "
    "span matching a known alias collapses to the group's canonical name, so that surface "
    "variants of one organisation aggregate under a single key. Validation checks "
    "plausibility: when an actor's recorded country of operation contradicts the country "
    "implied by the location extracted from the same sentence — the M23 operating in "
    "Maiduguri, say — the event's aggregate confidence is lowered and the record is "
    "flagged for human review. The knowledge base thus functions as an audit layer over "
    "the learned extractor rather than as a component of it.")
add_table(
    "Knowledge-base contents.",
    ["Resource", "Approximate size"],
    [
        ["Armed groups (with aliases, country, region, type)", "150"],
        ["Conflict-affected cities (mapped to country and admin unit)", "200"],
        ["African countries with primary regions", "54"],
        ["Weapon types and tactical-method categories", "30"],
        ["Taxonomy terminal categories", "95"],
    ],
    col_widths=[11.0, 4.9],
)

# ================= 4 EXPERIMENTAL SETUP =================
heading("4. Experimental Setup", 1)
para(
    "All results are reported on the 10,000-example validation partition, which was carved "
    "off before training began and used only to select the best checkpoint — never for "
    "hyperparameter tuning or sampler design. Evaluation is span-level under strict "
    "matching as implemented in seqeval [26]: a predicted span counts as correct only when "
    "its type, start, and end all coincide with a gold span, and a partial overlap counts "
    "as both a false positive and a false negative. Strict scoring was chosen because the "
    "downstream consumers of an extracted record — the validator, the event store, "
    "aggregation queries — treat a boundary-mismatched span as a different extraction, "
    "not a partial one. Both macro-averaged F1 (each entity type weighted equally, the "
    "right figure for assessing balance) and micro-averaged F1 (counts pooled across "
    "types, the right figure for overall throughput) are reported. Table 5 shows the "
    "entity distribution of the underlying data, which motivates the entire "
    "imbalance-handling apparatus: the five frequent types account for over 84% of entity "
    "tokens, while the outside label alone covers roughly 78% of all tokens.")
add_table(
    "Entity distribution in the pre-processed ACLED-derived corpus.",
    ["Entity", "Token count", "Share of entity tokens"],
    [
        ["ACTOR", "242,302", "23.9%"],
        ["CITY", "222,065", "21.9%"],
        ["DATE", "160,664", "15.9%"],
        ["REGION", "121,892", "12.0%"],
        ["DISTRICT", "107,453", "10.6%"],
        ["ACTION", "50,109", "4.9%"],
        ["VICTIM", "27,641", "2.7%"],
        ["CASUALTIES", "24,634", "2.4%"],
    ],
    col_widths=[5.0, 5.0, 5.9],
)
para(
    "For the loss-function ablation, four models were trained under identical conditions "
    "— same data, scheduler, early stopping, and random seeds — varying only the "
    "objective: plain cross-entropy, class-weighted cross-entropy, focal loss alone "
    "(γ = 2), and focal loss with inverse-frequency weights. This isolates the "
    "contribution of each imbalance treatment and of their combination.")

# ================= 5 RESULTS =================
heading("5. Results", 1)
heading("5.1 Convergence and loss-function comparison", 2)
para(
    "Across every configuration the model converges fast: validation loss reaches its "
    "minimum at epoch 2 for the focal configurations and creeps upward thereafter while "
    "training loss continues to fall — textbook overfitting, terminated by early "
    "stopping. Rapid convergence is expected when the encoder arrives pre-trained and the "
    "fine-tuning corpus saturates quickly. One detail deserves note: token-level validation "
    "accuracy keeps rising for several epochs after validation loss has begun to "
    "deteriorate. There is no contradiction — the focal objective makes the model more "
    "confident on tokens it already classifies correctly (raising accuracy) while its "
    "calibration on the minority-class boundaries it still gets wrong degrades (raising "
    "loss). Checkpoint selection therefore trusts the loss, not the accuracy. Table 6 "
    "compares the loss configurations; each ingredient that addresses imbalance lowers the "
    "best validation loss further, and the combination is best.")
add_table(
    "Best validation metrics by training objective.",
    ["Objective", "Best epoch", "Validation loss", "Token accuracy"],
    [
        ["Cross-entropy (baseline)", "4", "0.0102", "95.8%"],
        ["Class-weighted cross-entropy", "3", "0.0085", "96.4%"],
        ["Focal loss (γ = 2)", "2", "0.0079", "96.7%"],
        ["Focal + class weights", "2", "0.0074", "96.7%"],
        ["Focal + weights + smoothing (β = 0.1)", "2", "0.0076", "96.6%"],
    ],
    col_widths=[7.0, 2.6, 3.4, 3.4],
)
heading("5.2 Per-entity performance", 2)
para(
    "Table 7 reports span-level precision, recall, and F1 per entity for the production "
    "configuration. The model reaches macro F1 of 0.887 and micro F1 of 0.909. Three tiers "
    "are visible. Dates are easiest (F1 0.956): temporal expressions follow few patterns "
    "and are plentiful. Actors and cities form a strong second cluster, backed by rich "
    "training distributions and distinctive surface forms. The trailing entities are "
    "districts (0.826), which lose most of their errors to mutual confusion with cities "
    "and regions — many African place names legitimately denote both a town and the "
    "administrative unit around it — and victims (0.817), at once the rarest class and "
    "the one with the most variable phrasing. That every entity type, including the "
    "rarest, exceeds 0.8 is the operationally meaningful property: the macro figure "
    "certifies balance, while the micro figure, dominated by the high-support types, "
    "estimates overall extraction quality.")
add_table(
    "Per-entity precision, recall, and F1 on the held-out validation set (strict span matching).",
    ["Entity", "Support", "Precision", "Recall", "F1"],
    [
        ["ACTOR", "47,612", "0.929", "0.917", "0.923"],
        ["CITY", "44,361", "0.941", "0.928", "0.934"],
        ["DATE", "31,938", "0.961", "0.952", "0.956"],
        ["REGION", "24,331", "0.902", "0.881", "0.891"],
        ["DISTRICT", "21,471", "0.842", "0.811", "0.826"],
        ["ACTION", "9,963", "0.881", "0.852", "0.866"],
        ["VICTIM", "5,492", "0.838", "0.798", "0.817"],
        ["CASUALTIES", "4,907", "0.901", "0.869", "0.885"],
        ["Macro average", "—", "0.899", "0.876", "0.887"],
        ["Micro average", "190,075", "0.918", "0.901", "0.909"],
    ],
    col_widths=[4.4, 2.9, 2.9, 2.9, 2.8],
)
heading("5.3 Ablation: are the two imbalance treatments complementary?", 2)
para(
    "Table 8 breaks the ablation down per entity. The pattern is unambiguous. The minority "
    "entities benefit most from the combination: relative to plain cross-entropy, victim "
    "F1 rises by 10.9 points and action by 7.2. Neither ingredient carries the other — "
    "class weighting alone buys victims roughly seven points and focal loss alone nine, "
    "while the pair together buys eleven — so the two mechanisms are complementary "
    "rather than redundant, consistent with their acting on different axes (class "
    "frequency versus example difficulty). Just as important, no entity is harmed: the "
    "high-support types show small consistent gains. An imbalance treatment that protected "
    "rare classes by sacrificing the frequent ones would be a poor trade for analytical "
    "workflows that depend on reliable actor and date extraction; this one makes no such "
    "trade.")
add_table(
    "Per-entity F1 across the loss-function ablation.",
    ["Entity", "Cross-entropy", "Weighted CE", "Focal (γ = 2)", "Focal + weights"],
    [
        ["ACTOR", "0.914", "0.918", "0.920", "0.923"],
        ["CITY", "0.929", "0.931", "0.931", "0.934"],
        ["DATE", "0.953", "0.955", "0.955", "0.956"],
        ["REGION", "0.879", "0.884", "0.887", "0.891"],
        ["DISTRICT", "0.808", "0.815", "0.821", "0.826"],
        ["ACTION", "0.794", "0.834", "0.842", "0.866"],
        ["VICTIM", "0.708", "0.776", "0.792", "0.817"],
        ["CASUALTIES", "0.853", "0.871", "0.872", "0.885"],
        ["Macro average", "0.855", "0.873", "0.878", "0.887"],
    ],
    col_widths=[4.0, 3.0, 3.0, 3.0, 3.0],
)
heading("5.4 Impact of the knowledge-base layer", 2)
para(
    "The knowledge base operates downstream of the extractor and does not move F1, so its "
    "contribution is measured operationally, one metric per function. On the enrichment "
    "side, 64.3% of high-confidence actor spans (confidence ≥ 0.85) on the validation "
    "set match a canonical knowledge-base entry and therefore reach the analyst with "
    "country, region, and group-type metadata already attached; the remaining 35.7% are "
    "overwhelmingly generic descriptors (“gunmen”, “armed men”, “the "
    "attackers”) that no reference resource could canonicalise without additional "
    "context. Roughly two-thirds of named perpetrators thus aggregate automatically under "
    "a single identifier — the property that makes actor-level queries trustworthy. On "
    "the validation side, the geographic-plausibility check fires on 2.4% of multi-entity "
    "events whose extracted city and region imply different countries. Manual inspection "
    "shows the flagged cases are mostly genuine cross-border incidents or extraction "
    "errors at the region–city boundary; in both cases the flag correctly surfaces a "
    "record a human should re-read before trusting. Neither figure is dramatic, and that "
    "is the point: the layer adds metadata to most records and asks for human attention "
    "on one in forty, without filtering away valid extractions.")
heading("5.5 Error analysis", 2)
para(
    "Three hundred validation events on which the model made at least one error were read "
    "individually. Five patterns account for essentially all of them. Boundary mismatches "
    "dominate at roughly 38%: the model finds the right entity but clips a qualifier "
    "— “at least 12 civilians” becomes “12 civilians” — and "
    "victims and casualties absorb nearly all of these. Strict scoring counts each clipped "
    "span as a complete miss, so the headline F1 understates the usefulness of the output "
    "an analyst actually sees. Confusion among the three location types contributes about "
    "25%, concentrated on places that are genuinely two things at once (Goma the city "
    "versus Goma the seat of North Kivu); the model defaults to the city reading, which "
    "is right more often than wrong but guarantees a steady trickle of strict-scoring "
    "errors. Outright misses make up 19%, concentrated on victim phrasings the "
    "augmentation templates never covered and on passive-voice action verbs. Spurious "
    "extractions run at 12%, dominated by vague temporal expressions (“this "
    "morning”, “earlier”) tagged as dates; raising the date threshold to 0.85 "
    "would remove most of them at a measured cost of about 1.2 points of legitimate date "
    "recall, a dial left to the operator. The final 7% are correct predictions whose "
    "averaged confidence lands just below the category threshold. Each pattern maps to a "
    "concrete remedy — span-boundary modelling, knowledge-injected disambiguation of "
    "location types, and negative examples for vague dates — taken up in Section 7.")

# ================= 6 DISCUSSION =================
heading("6. Discussion", 1)
para(
    "The single most consequential design decision was not architectural but definitional: "
    "restricting supervision to entities that can be grounded verbatim in text. The point "
    "is best made by the failure that preceded it. An early attempt to learn event type as "
    "a first-class NER label — on a schema where only about 58% of its annotations "
    "could be located in the source — plateaued near F1 0.4 and degraded neighbouring "
    "entities, because the model was being asked to reproduce annotator inferences rather "
    "than recognise textual evidence. Moving event type and country out of the supervised "
    "problem and into deterministic post-processing recovered both cheaply and let the "
    "eight grounded entities train cleanly. The general lesson for applied NER schema "
    "design is to treat groundability as a measurable property to be piloted, not assumed: "
    "the question is not whether a category is analytically useful but whether its "
    "evidence is in the text.")
para(
    "The second finding with likely generality is the complementarity of focal loss and "
    "class weighting. Eleven F1 points on the rarest class from what is essentially a "
    "single change to the training objective is a large effect, and the ablation shows it "
    "is not attributable to either ingredient alone. Since the two mechanisms address "
    "different failure modes — weighting corrects for how rarely a class is seen, "
    "focusing corrects for how easily it is learned — there is little reason to choose "
    "between them in other severely imbalanced token-classification settings, and some "
    "evidence here that choosing between them leaves accuracy on the table.")
para(
    "Third, the knowledge base earns its keep at the operational end of the pipeline "
    "rather than the modelling end. Its measurable contributions — canonicalising "
    "two-thirds of named actors, flagging one record in forty for review — do not "
    "appear in any F1 column, yet they bear directly on whether aggregate queries over "
    "extracted data can be trusted, which is the property that determines adoption. "
    "Evaluations of extraction systems intended for operational use should report such "
    "downstream measures alongside intrinsic accuracy.")
para(
    "Several limitations qualify the results. The reported figures are single-run point "
    "estimates; informal repetition during development suggests macro F1 varies by "
    "roughly ±0.5 points across seeds, but a formal multi-seed study remains to be "
    "done. The validation partition is drawn from the same combined corpus as training, "
    "so roughly 30% of it is synthetic; the estimates are therefore fair in-distribution "
    "but optimistic for out-of-distribution prose, and evaluation on a fully natural, "
    "independently annotated news sample is the most important pending experiment. The "
    "corpus inherits ACLED's concise editorial style, and generalisation to long-form "
    "journalism, machine-translated articles, or citizen-journalist text is unmeasured. "
    "The system is bound to English, while a large share of African conflict reporting is "
    "in French, Arabic, Portuguese, and African languages. Finally, the model performs "
    "flat NER; nested and overlapping mentions are not represented, and the 5W1H frame "
    "deliberately excludes motive and consequence, which are inference problems rather "
    "than extraction problems.")

# ================= 7 CONCLUSION =================
heading("7. Conclusion", 1)
para(
    "This paper addressed the conversion bottleneck between African conflict reporting and "
    "the structured event records that monitoring and research depend on. The approach "
    "rests on three commitments: supervise only what the text grounds, treat severe class "
    "imbalance in the objective rather than hoping scale absorbs it, and audit learned "
    "output against curated domain knowledge. A BERT model fine-tuned under focal loss "
    "with inverse-frequency class weighting on a 50,000-example ACLED-derived corpus "
    "reaches macro F1 of 0.887 and micro F1 of 0.909 over an eight-entity 5W1H schema, "
    "with the controlled ablation attributing an eleven-point gain on the rarest entity "
    "to the combined objective. The released taxonomy, annotation guidelines, and "
    "knowledge base are intended as standalone resources for conflict-monitoring and "
    "African NLP research beyond this system.")
para(
    "Future work follows directly from the error analysis and the limitations above. The "
    "highest priorities are an independently annotated natural-news evaluation set and a "
    "formal multi-seed study to firm up the estimates; multilingual extension through an "
    "African-pre-trained or multilingual backbone such as AfroLM or XLM-RoBERTa; explicit "
    "span-boundary modelling (a CRF or biaffine span classifier over the encoder) to "
    "attack the dominant boundary-error class; replacement of the rule-based taxonomy "
    "step with a learned hierarchical classifier; and an active-learning loop that routes "
    "low-confidence and knowledge-base-flagged extractions to human review and feeds the "
    "corrections back into training.")

# ================= BACK MATTER =================
heading("Data Availability", 1)
para(
    "The hierarchical taxonomy of African violent events, the entity annotation "
    "guidelines, and the curated knowledge base of armed groups and conflict-affected "
    "locations are publicly available at [repository URL to be inserted upon "
    "acceptance]. The training corpus is derived from ACLED data and is therefore "
    "available from the authors on reasonable request, subject to the ACLED terms of "
    "use; researchers may reconstruct it from the ACLED open-data export using the "
    "documented procedure in Section 3.3.")
heading("Acknowledgements", 1)
para(
    "[Acknowledge advisor, institution, and any colleagues or analysts who participated "
    "in evaluation.]")
heading("Declarations", 1)
para("Conflict of interest: The authors declare no conflict of interest.", space_after=2)
para("Funding: This research received no external funding.", space_after=2)
para(
    "Note: This article is derived from the first author's master's thesis at Addis "
    "Ababa University; the thesis has not been published in a journal or conference "
    "venue.", space_after=10)

heading("References", 1)
REFS = [
    "C. Raleigh, A. Linke, H. Hegre, and J. Karlsen, “Introducing ACLED: An armed conflict location and event dataset,” Journal of Peace Research, vol. 47, no. 5, pp. 651–660, 2010.",
    "R. Sundberg and E. Melander, “Introducing the UCDP georeferenced event dataset,” Journal of Peace Research, vol. 50, no. 4, pp. 523–532, 2013.",
    "K. Leetaru and P. A. Schrodt, “GDELT: Global data on events, location, and tone, 1979–2012,” in ISA Annual Convention, vol. 2, no. 4, 2013, pp. 1–49.",
    "A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, L. Kaiser, and I. Polosukhin, “Attention is all you need,” in Advances in Neural Information Processing Systems, vol. 30, 2017.",
    "J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, “BERT: Pre-training of deep bidirectional transformers for language understanding,” in Proceedings of NAACL-HLT, Minneapolis, MN, 2019, pp. 4171–4186.",
    "D. I. Adelani et al., “MasakhaNER: Named entity recognition for African languages,” Transactions of the Association for Computational Linguistics, vol. 9, pp. 1116–1131, 2021.",
    "G. R. Doddington, A. Mitchell, M. A. Przybocki, L. A. Ramshaw, S. M. Strassel, and R. M. Weischedel, “The Automatic Content Extraction (ACE) program: Tasks, data, and evaluation,” in Proceedings of the 4th International Conference on Language Resources and Evaluation, Lisbon, Portugal, 2004, pp. 837–840.",
    "D. Ahn, “The stages of event extraction,” in Proceedings of the Workshop on Annotating and Reasoning about Time and Events, Sydney, Australia, 2006, pp. 1–8.",
    "M. Liu, B. Liu, L. Liu, M. Wang, and X. Zhou, “Event extraction as machine reading comprehension,” in Proceedings of EMNLP, 2020, pp. 1641–1651.",
    "F. Hogenboom, F. Frasincar, U. Kaymak, F. de Jong, and E. Caron, “A survey of event extraction methods from text for decision support systems,” Decision Support Systems, vol. 85, pp. 12–22, 2016.",
    "H. Tanev, M. Atkinson, and J. Piskorski, “Real-time news event extraction for global crisis monitoring,” in 13th International Conference on Natural Language and Information Systems, LNCS, vol. 5039, 2008, pp. 207–218.",
    "J. Piskorski, H. Tanev, and P. O. Wennerberg, “Extracting violent events from on-line news for ontology population,” in Business Information Systems, LNCS, vol. 4439, 2007, pp. 287–300.",
    "J. Lafferty, A. McCallum, and F. C. N. Pereira, “Conditional random fields: Probabilistic models for segmenting and labeling sequence data,” in Proceedings of the 18th International Conference on Machine Learning, 2001, pp. 282–289.",
    "G. Lample, M. Ballesteros, S. Subramanian, K. Kawakami, and C. Dyer, “Neural architectures for named entity recognition,” in Proceedings of NAACL-HLT, San Diego, CA, 2016, pp. 260–270.",
    "X. Ma and E. Hovy, “End-to-end sequence labeling via bi-directional LSTM-CNNs-CRF,” in Proceedings of the 54th Annual Meeting of the Association for Computational Linguistics, Berlin, Germany, 2016, pp. 1064–1074.",
    "B. F. P. Dossou et al., “AfroLM: A self-active learning-based multilingual pretrained language model for 23 African languages,” in Proceedings of the 3rd Workshop on Simple and Efficient Natural Language Processing (SustaiNLP) at EMNLP, Abu Dhabi, UAE, 2022, pp. 52–64.",
    "A. Conneau et al., “Unsupervised cross-lingual representation learning at scale,” in Proceedings of ACL, 2020, pp. 8440–8451.",
    "N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, “SMOTE: Synthetic minority over-sampling technique,” Journal of Artificial Intelligence Research, vol. 16, pp. 321–357, 2002.",
    "H. He and E. A. Garcia, “Learning from imbalanced data,” IEEE Transactions on Knowledge and Data Engineering, vol. 21, no. 9, pp. 1263–1284, 2009.",
    "Y. Cui, M. Jia, T.-Y. Lin, Y. Song, and S. Belongie, “Class-balanced loss based on effective number of samples,” in Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 2019, pp. 9268–9277.",
    "T.-Y. Lin, P. Goyal, R. Girshick, K. He, and P. Dollár, “Focal loss for dense object detection,” in Proceedings of the IEEE International Conference on Computer Vision, 2017, pp. 2980–2988.",
    "D. J. Gerner, P. A. Schrodt, O. Yilmaz, and R. Abu-Jabr, “Conflict and Mediation Event Observations (CAMEO): A new event data framework for the analysis of foreign policy interactions,” in International Studies Association Annual Meeting, New Orleans, LA, 2002.",
    "T. A. Edris and R. K. Sungkur, “Knowledge discovery from free text: Extraction of violent events in the African context,” New Review of Information Networking, vol. 24, no. 2, pp. 153–177, 2019.",
    "W. Wang and D. Zhao, “Chinese news event 5W1H semantic elements extraction for event ontology population,” in Proceedings of the 21st International World Wide Web Conference (Companion Volume), Lyon, France, 2012, pp. 197–202.",
    "C. N. Silla and A. A. Freitas, “A survey of hierarchical classification across different application domains,” Data Mining and Knowledge Discovery, vol. 22, no. 1–2, pp. 31–72, 2011.",
    "H. Nakayama, “seqeval: A Python framework for sequence labeling evaluation,” Software, 2018. [Online]. Available: https://github.com/chakki-works/seqeval",
]
for i, ref in enumerate(REFS, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.first_line_indent = Cm(-0.9)
    r = p.add_run("[%d] %s" % (i, ref))
    r.font.size = Pt(11)

doc.save(OUT)
print("Saved:", OUT)
