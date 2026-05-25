#!/usr/bin/env python3
"""Build thesis.docx from thesis.md with AAU CS formatting.

Pipeline:
  1. Generate a reference.docx that pins A4 paper, 1.3"/1" margins,
     Times New Roman 12 pt, 1.5 line spacing, and justified body
     paragraphs. Heading 1/2/3 inherit Times New Roman in graduated
     sizes, bold, left-aligned.
  2. Copy thesis.md to a temp file. Strip the static Table of Contents
     block (HTML comment + heading + table) from where it currently
     sits (after Acknowledgements), and re-insert just the heading
     plus a TOC sentinel right after the signature page (before the
     Abstract). The static List of Tables / Figures / Algorithms
     tables are preserved in place.
  3. Invoke pandoc with the reference doc, producing a docx whose
     headings are real Word Heading styles.
  4. Post-process:
       - Replace the sentinel paragraph with a real Word TOC field
         scoped by a bookmark range, so the TOC lists Abstract
         onward and excludes the title page and the TOC heading
         itself.
       - Add a 0.5 pt single-line border to every table.
       - Clear any direct alignment overrides on body paragraphs so
         the justified Normal style wins everywhere.
       - Set w:updateFields=true so Word silently repaginates the
         TOC on first open.

Source thesis.md is never mutated; all transforms run on a temp copy.

Usage:
    cd backend/docs/thesis
    python build_thesis.py
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HERE = Path(__file__).resolve().parent
SRC_MD = HERE / "thesis.md"
OUT_DOCX = HERE / "thesis.docx"
REFERENCE_DOCX = HERE / "reference.docx"

# A4 paper. Pandoc uses the reference doc's section properties verbatim.
A4_WIDTH = Inches(8.27)
A4_HEIGHT = Inches(11.69)
MARGIN_LEFT = Inches(1.3)
MARGIN_OTHER = Inches(1.0)

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
LINE_SPACING = 1.5

# Styles pandoc emits for body content. All three should follow the
# Normal typography rules so the body looks uniform regardless of which
# style pandoc picked for a given paragraph.
BODY_STYLE_NAMES = ("Normal", "Body Text", "First Paragraph", "Compact")

# Pandoc tags fenced code blocks with the "Source Code" paragraph style.
# It needs a monospaced font, a smaller size, left alignment, and tight
# line spacing — otherwise the body's justified Times New Roman ruins
# any ASCII art or column-aligned pseudocode.
CODE_FONT = "Courier New"
CODE_SIZE = Pt(10)
CODE_LINE_SPACING = 1.0

HEADING_SIZES = {
    "Heading 1": Pt(16),
    "Heading 2": Pt(14),
    "Heading 3": Pt(13),
    "Heading 4": Pt(12),
    "Heading 5": Pt(12),
    "Heading 6": Pt(12),
}


def configure_body_style(style) -> None:
    font = style.font
    font.name = BODY_FONT
    font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), BODY_FONT)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_code_style(style) -> None:
    font = style.font
    font.name = CODE_FONT
    font.size = CODE_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), CODE_FONT)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = CODE_LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)


def configure_heading_style(style, size) -> None:
    font = style.font
    font.name = BODY_FONT
    font.size = size
    font.bold = True
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), BODY_FONT)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.keep_with_next = True


def build_reference_docx(path: Path) -> None:
    """Programmatically build the pandoc reference template."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.top_margin = MARGIN_OTHER
    section.right_margin = MARGIN_OTHER
    section.bottom_margin = MARGIN_OTHER

    styles = doc.styles
    for name in BODY_STYLE_NAMES:
        try:
            configure_body_style(styles[name])
        except KeyError:
            # Style not present in the template; pandoc will create it
            # at write time and inherit from Normal.
            pass

    for name, size in HEADING_SIZES.items():
        try:
            configure_heading_style(styles[name], size)
        except KeyError:
            pass

    # Source Code style for fenced code blocks (algorithms, ASCII tree,
    # etc.). python-docx ships only the styles the bundled template
    # already defines, so we create "Source Code" if it is missing.
    try:
        sc_style = styles["Source Code"]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE

        sc_style = styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    configure_code_style(sc_style)

    # Caption + List Paragraph + Quote should follow body typography.
    for name in ("Caption", "List Paragraph", "Quote", "Intense Quote"):
        try:
            s = styles[name]
            s.font.name = BODY_FONT
            s.font.size = BODY_SIZE
            s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            s.paragraph_format.line_spacing = LINE_SPACING
            if name in ("List Paragraph", "Quote", "Intense Quote"):
                s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass

    # Do NOT set <w:updateFields w:val="true"/>. With it set, Word's
    # default Trust Center configuration intercepts the auto-update and
    # shows the dialog: "This document contains fields that may refer
    # to other files. Do you want to update the fields?" — which is the
    # opposite of what we want.
    #
    # Instead we leave the flag unset. The TOC field still ships in the
    # document; the user populates it once via right-click → Update
    # Field, saves, and from then on every open is prompt-free with the
    # TOC already filled in.
    settings_el = doc.settings.element
    existing = settings_el.find(qn("w:updateFields"))
    if existing is not None:
        settings_el.remove(existing)

    doc.save(path)


TOC_SENTINEL = "TOCFIELDANCHOR0001"
TOC_BOOKMARK = "TOC_RANGE"

# Matches the whole static TOC block including the leading HTML comment,
# the page break that precedes it, and the trailing page break before the
# List of Tables. We delete this block and re-insert the heading
# elsewhere; the trailing page break is reinstated so Acknowledgements
# still gets a page break before List of Tables.
OLD_TOC_BLOCK_RE = re.compile(
    r"\n<!--\s*\nStatic Table of Contents.*?"
    r"\|\s*Signed Declaration Sheet[^\n]*\n\s*\n\\pagebreak\s*\n\s*\n"
    r"# List of Tables",
    re.DOTALL,
)

# Anchor for inserting the new TOC heading. "# Abstract" is unique in the
# document.
ABSTRACT_HEADING_RE = re.compile(r"\\pagebreak\s*\n\s*\n# Abstract\b")


def reorder_for_build(md_text: str) -> str:
    """Strip the old TOC block and insert the heading + sentinel before
    the Abstract section."""
    new_text, n = OLD_TOC_BLOCK_RE.subn(
        lambda _m: "\n\\pagebreak\n\n# List of Tables", md_text, count=1
    )
    if n != 1:
        raise RuntimeError(
            "Could not locate the static Table of Contents block in thesis.md"
        )

    toc_insert = (
        "\\pagebreak\n\n"
        "# Table of Contents\n\n"
        f"{TOC_SENTINEL}\n\n"
        "\\pagebreak\n\n"
        "# Abstract"
    )
    new_text, n = ABSTRACT_HEADING_RE.subn(
        lambda _m: toc_insert, new_text, count=1
    )
    if n != 1:
        raise RuntimeError(
            "Could not locate the Abstract heading in thesis.md"
        )
    return new_text


def run_pandoc(src_md: Path, out_docx: Path, reference: Path) -> None:
    cmd = [
        "pandoc",
        str(src_md),
        "-o",
        str(out_docx),
        f"--reference-doc={reference}",
    ]
    print(" ".join(cmd))
    subprocess.run(cmd, check=True)


def _new_run() -> OxmlElement:
    return OxmlElement("w:r")


def _fldChar(kind: str) -> OxmlElement:
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), kind)
    return el


def _instr(text: str) -> OxmlElement:
    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = text
    return el


def _text(text: str) -> OxmlElement:
    el = OxmlElement("w:t")
    el.set(qn("xml:space"), "preserve")
    el.text = text
    return el


def insert_toc_field(paragraph, bookmark_name: str = TOC_BOOKMARK) -> None:
    """Replace the contents of *paragraph* with a Word TOC field.

    The field is constructed so Word renders it as a navigable Table of
    Contents covering Heading 1-3 with hyperlinked entries, scoped to
    *bookmark_name* so the title page and the TOC heading itself are
    excluded.
    """
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)

    begin = _new_run()
    begin.append(_fldChar("begin"))
    p.append(begin)

    instr_run = _new_run()
    instr_run.append(
        _instr(f' TOC \\o "1-3" \\h \\z \\u \\b "{bookmark_name}" ')
    )
    p.append(instr_run)

    sep = _new_run()
    sep.append(_fldChar("separate"))
    p.append(sep)

    placeholder = _new_run()
    placeholder.append(
        _text(
            "Right-click this line and choose “Update Field” "
            "to populate the Table of Contents."
        )
    )
    p.append(placeholder)

    end = _new_run()
    end.append(_fldChar("end"))
    p.append(end)


def find_toc_anchor(doc):
    """Find the paragraph that contains the <!--TOC_FIELD--> marker.

    Pandoc emits the HTML comment as a literal text run in a paragraph
    immediately after the "Table of Contents" heading.
    """
    for para in doc.paragraphs:
        if TOC_SENTINEL in para.text:
            return para
    raise RuntimeError(f"{TOC_SENTINEL} anchor not found in generated docx")


def enforce_body_alignment(doc) -> None:
    """Belt-and-braces: force body-style paragraphs to justified.

    Pandoc occasionally writes paragraphs with a direct alignment hint
    that overrides the Normal style. We strip those so the style wins,
    which keeps everything visually uniform.
    """
    body_styles = {
        "Normal",
        "Body Text",
        "First Paragraph",
        "Compact",
        "List Paragraph",
    }
    for para in doc.paragraphs:
        try:
            style_name = para.style.name
        except AttributeError:
            continue
        if style_name not in body_styles:
            continue
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            pPr.remove(jc)


def add_toc_bookmark(doc, bookmark_name: str = TOC_BOOKMARK) -> None:
    """Wrap a Word bookmark from the Abstract heading to the end of the
    body so the TOC field can be scoped to only that range.

    Scoping the TOC excludes the title-page H1 (which appears before
    the TOC heading) and the TOC heading itself from the generated
    entries.
    """
    abstract_para = None
    for para in doc.paragraphs:
        if (
            para.text.strip() == "Abstract"
            and para.style.name == "Heading 1"
        ):
            abstract_para = para
            break
    if abstract_para is None:
        raise RuntimeError(
            "Could not find Abstract heading to anchor TOC bookmark"
        )

    bm_id = "1000"
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), bm_id)
    bm_start.set(qn("w:name"), bookmark_name)
    pPr = abstract_para._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bm_start)
    else:
        abstract_para._p.insert(0, bm_start)

    last_para = doc.paragraphs[-1]
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), bm_id)
    last_para._p.append(bm_end)


def apply_table_borders(doc) -> None:
    """Apply a 0.5 pt single-line black border to every cell of every
    table. Pandoc emits tables with no borders by default, which the
    AAU CS guideline does not accept."""
    border_attrs = {
        "w:val": "single",
        "w:sz": "4",  # eighths of a point → 0.5 pt
        "w:space": "0",
        "w:color": "000000",
    }
    edges = ("top", "left", "bottom", "right", "insideH", "insideV")
    cell_edges = ("top", "left", "bottom", "right")

    for table in doc.tables:
        tbl = table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tbl_borders = OxmlElement("w:tblBorders")
        for edge in edges:
            el = OxmlElement(f"w:{edge}")
            for k, v in border_attrs.items():
                el.set(qn(k), v)
            tbl_borders.append(el)
        tblPr.append(tbl_borders)

        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                old_cell_borders = tcPr.find(qn("w:tcBorders"))
                if old_cell_borders is not None:
                    tcPr.remove(old_cell_borders)
                tc_borders = OxmlElement("w:tcBorders")
                for edge in cell_edges:
                    el = OxmlElement(f"w:{edge}")
                    for k, v in border_attrs.items():
                        el.set(qn(k), v)
                    tc_borders.append(el)
                tcPr.append(tc_borders)


def post_process(out_docx: Path) -> None:
    doc = Document(out_docx)
    anchor = find_toc_anchor(doc)
    add_toc_bookmark(doc)
    insert_toc_field(anchor)
    apply_table_borders(doc)
    enforce_body_alignment(doc)
    doc.save(out_docx)


def main() -> int:
    if not SRC_MD.exists():
        print(f"thesis.md not found at {SRC_MD}", file=sys.stderr)
        return 1
    if shutil.which("pandoc") is None:
        print("pandoc is not on PATH", file=sys.stderr)
        return 1

    print(f"Building reference template at {REFERENCE_DOCX}")
    build_reference_docx(REFERENCE_DOCX)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_md = Path(tmp) / "thesis.md"
        md_text = SRC_MD.read_text(encoding="utf-8")
        tmp_md.write_text(reorder_for_build(md_text), encoding="utf-8")
        run_pandoc(tmp_md, OUT_DOCX, REFERENCE_DOCX)

    print(f"Post-processing {OUT_DOCX}")
    post_process(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
